"""
下载KM文章图片并创建Word文档
"""
import os
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# 设置工作目录
WORK_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(WORK_DIR, 'images')
os.makedirs(IMAGES_DIR, exist_ok=True)

# 文章图片URL列表
IMAGE_URLS = [
    'https://km.woa.com/asset/00010002251200f63a486c1c4843b902?height=1120&width=1658&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200ba522e54f45642d402?height=1350&width=1040&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/0001000225120097ffb87952ff46c602?height=688&width=1180&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200ab0d49a2271e4ebc02?height=953&width=1170&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200a505422be550439d02?height=440&width=1183&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200414c928523de45db02?height=1410&width=1200&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200b65abd960591448a02?height=630&width=1190&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200de697fc7fa6b4c7d02?height=1280&width=3395&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
    'https://km.woa.com/asset/00010002251200a1438deb9654432d02?height=248&width=1005&imageMogr2/thumbnail/1540x%3E/ignore-error/1',
]

# 图片描述
IMAGE_DESCRIPTIONS = [
    '图1：出海合规智能决策系统架构图',
    '图2：问题根因分析（5Why分析法）',
    '图3：传统流程 vs AI赋能流程对比',
    '图4：AI Agent与知识库落地方案',
    '图5：工作流规划示意图',
    '图6：场景化工作流详细构建方案',
    '图7：知识库构建关键技术点',
    '图8：系统综合应用界面',
    '图9：运营监控体系',
]

def download_images(session):
    """下载所有图片"""
    image_paths = []
    for i, url in enumerate(IMAGE_URLS):
        try:
            print(f"下载图片 {i+1}/{len(IMAGE_URLS)}...")
            resp = session.get(url, timeout=30)
            resp.raise_for_status()
            
            # 保存图片
            img_path = os.path.join(IMAGES_DIR, f'image_{i+1}.png')
            with open(img_path, 'wb') as f:
                f.write(resp.content)
            image_paths.append(img_path)
            print(f"  已保存: {img_path}")
        except Exception as e:
            print(f"  下载失败: {e}")
            image_paths.append(None)
    return image_paths

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document(image_paths):
    """创建Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    
    # ========== 标题 ==========
    title = doc.add_heading('', level=0)
    title_run = title.add_run('扬帆出海：AI赋能云业务全球化合规实战')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('（含技术实现核心方案）')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 作者信息
    author_para = doc.add_paragraph()
    author_run = author_para.add_run('作者：frankechen  |  发布时间：2025-11-30')
    author_run.font.size = Pt(10)
    author_run.font.color.rgb = RGBColor(128, 128, 128)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # ========== 导语 ==========
    doc.add_heading('导语', level=1)
    doc.add_paragraph(
        '从0到1构建智能出海合规助手的完整方法论，基于腾讯云真实出海项目经验沉淀和AI平台赋能。'
    )
    doc.add_paragraph(
        '核心问题：企业出海决策周期长(6+个月)、成本高、重复劳动严重等。'
    )
    doc.add_paragraph(
        '解决方案：通过AI+知识库，打造智能方案和合规审核系统。'
    )
    
    # 插入第一张图片
    if image_paths[0]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[0], width=Inches(6))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[0])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # ========== 第一部分：背景篇 ==========
    doc.add_heading('第一部分：背景篇', level=1)
    
    doc.add_heading('1.1 问题定义：出海三大核心痛点', level=2)
    
    doc.add_heading('痛点1：产品可售性黑盒', level=3)
    doc.add_paragraph('📍 真实场景还原')
    doc.add_paragraph('时间：202X年X月 | 地点：腾讯云深圳总部 | 人物：销售XX vs 客户XX集团')
    
    # 对话框
    dialog = doc.add_paragraph()
    dialog.paragraph_format.left_indent = Inches(0.5)
    dialog_text = '''客户: "我们要在XX建数据中心，X,000台服务器的需求。"
销售: "没问题！我们XXX产品可以满足。"
客户: "大概多少钱，什么时候能交付？"
销售: "这个...让我确认一下产品团队，3天内回复您。"
客户: "3天？我们下周就要决策了，今天就能给方案。"

[3天后]
产品: "XXX需要Anatel认证，巴西还没做，需要6个月。"
销售: (内心OS: 糟了，项目要黄了...)'''
    dialog_run = dialog.add_run(dialog_text)
    dialog_run.font.name = 'Consolas'
    dialog_run.font.size = Pt(10)
    
    # 插入第二张图片
    if image_paths[1]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[1], width=Inches(5))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[1])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('问题根因分析', level=4)
    
    # 表格1：5Why分析
    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    headers = ['层级', '为什么？', '根因']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E6F3FF')
    
    data = [
        ['Why 1', '为什么销售不知道？', '产品认证信息缺乏共享'],
        ['Why 2', '为什么信息未共享？', '信息散落在各个产品线和前端团队'],
        ['Why 3', '为什么信息散落？', '缺乏统一的分享平台'],
        ['Why 4', '为什么没有统一平台？', '没有意识到信息共享重要性'],
        ['Why 5', '为什么没有意识到？', '缺乏出海项目失败成本核算和复盘机制']
    ]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table1.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('痛点2：税务方案迷宫', level=3)
    doc.add_paragraph('💰 XXX项目税务对比')
    doc.add_paragraph('初版税负高达4X%，利润大头都被税务黑洞吃掉，对客报价缺乏竞争力。')
    
    # ========== 1.2 方案创新 ==========
    doc.add_heading('1.2 方案创新：AI如何重构出海决策流程', level=2)
    
    # 插入第三张图片
    if image_paths[2]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[2], width=Inches(6))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[2])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('传统流程 vs AI赋能流程对比', level=3)
    
    # 表格2：对比
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['对比维度', '传统模式（4周）', 'AI模式（30秒）', '提升幅度']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E6F3FF')
    
    data2 = [
        ['评估周期', '4周（产品3天+法务2周+税务1周）', '90秒即时回答', '99.7%↓'],
        ['单项目成本', 'X（人工+外包）', 'X（AI边际成本）', '98%↓'],
        ['准确率', 'X%（人工主观判断）', 'X%（基于真实案例）', '41.5%↑']
    ]
    for i, row_data in enumerate(data2):
        for j, cell_text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # ========== 第二部分：方法论篇 ==========
    doc.add_heading('第二部分：方法论篇', level=1)
    
    doc.add_heading('2.1 方案实现：AI Agent与知识库落地', level=2)
    
    # 插入第四张图片
    if image_paths[3]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[3], width=Inches(5.5))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[3])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('工作流规划', level=3)
    
    workflow_text = '''【第一层：整体规划】
企业出海合规智能决策系统
通过AI+知识库，将100+出海项目经验沉淀为可复用的决策系统

【第二层：三大支柱】
├─ 支柱1：产品可售性查询（即时回答）
│   - 为什么重要？避免盲目报价导致项目失败
│   - 怎么实现？数据中心开区表 + 认证清单 + 历史案例
│   - 效果如何？响应时间从3天→90秒
│
├─ 支柱2：税务优化方案（自动生成对比表）
│   - 为什么重要？税负差异可达X%影响成交
│   - 怎么实现？税率矩阵 + 交易结构优化 + 自动计算
│
└─ 支柱3：专家智能匹配（实时推荐）
    - 减少跨部门多次找人
    - 怎么实现？专家库 + 意图识别 + 协同工作流
    - 效果如何？对接时间从X天→即时推荐'''
    
    workflow_para = doc.add_paragraph()
    workflow_para.paragraph_format.left_indent = Inches(0.3)
    workflow_run = workflow_para.add_run(workflow_text)
    workflow_run.font.name = 'Consolas'
    workflow_run.font.size = Pt(9)
    
    # 插入第五张图片
    if image_paths[4]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[4], width=Inches(6))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[4])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 价值预估
    value_para = doc.add_paragraph()
    value_run = value_para.add_run('价值预估：扬帆出海AI助手：让出海决策从6+个月缩短到X个月')
    value_run.font.bold = True
    
    doc.add_paragraph('核心价值：构建智能出海合规助手，通过AI+知识库，提前识别风险和解决，缩短整体流程时间。')
    
    doc.add_heading('结合出海痛点，构建5大场景和工作流', level=3)
    
    # 表格3：功能模块
    table3 = doc.add_table(rows=7, cols=2)
    table3.style = 'Table Grid'
    for i, header in enumerate(['功能模块', '核心价值']):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E6F3FF')
    
    data3 = [
        ['知识库检索', '查询目标国家法规和合规要求、产品文档'],
        ['场景分析', '收集项目信息，输出方案'],
        ['产品可行性查询', '30秒判断能否售卖'],
        ['税务优化方案', '节税方案规划'],
        ['对接人匹配', '即时推荐产品专家'],
        ['完整报告生成', '基于场景规划中']
    ]
    for i, row_data in enumerate(data3):
        for j, cell_text in enumerate(row_data):
            table3.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # 插入第六张图片
    if image_paths[5]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[5], width=Inches(5.5))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[5])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('场景化工作流详细构建方案', level=3)
    doc.add_paragraph('将高频痛点场景转化为标准化工作流，实现智能化问题解决。')
    doc.add_paragraph('核心架构包含五大核心场景和智能化工作流引擎，通过意图识别、知识库检索、LLM大模型分析、代码解释器+插件、工作流执行五个关键环节实现落地。')
    
    # ========== 2.2 知识体系 ==========
    doc.add_heading('2.2 构建清晰高效的知识体系', level=2)
    
    # 插入第七张图片
    if image_paths[6]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[6], width=Inches(6))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[6])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('技术规格', level=3)
    doc.add_heading('工具清单', level=4)
    
    # 表格4：工具清单
    table4 = doc.add_table(rows=5, cols=4)
    table4.style = 'Table Grid'
    for i, header in enumerate(['工具', '文件名', '功能', '代码量']):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E6F3FF')
    
    data4 = [
        ['🕷️ 网页爬虫', 'web_scraper.py', '批量爬取网页，HTML转Markdown', '~350行'],
        ['📄 PDF解析器', 'pdf_parser.py', '解析PDF，提取文本表格目录', '~400行'],
        ['✨ 格式化器', 'formatter.py', '标准化格式，文本分片，质量评分', '~500行'],
        ['⚙️ 批量处理器', 'batch_processor.py', '整合所有工具，一键批量处理', '~350行']
    ]
    for i, row_data in enumerate(data4):
        for j, cell_text in enumerate(row_data):
            table4.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('知识库构建关键技术点', level=3)
    doc.add_paragraph('• 自动化搜集：基于沙箱环境的网页爬虫 + PDF解析，批量获取1600+篇文档，节省90%人工时间')
    doc.add_paragraph('• 智能格式化：标准Markdown格式，7维度质量评分，确保文档质量')
    doc.add_paragraph('• RAG优化：针对BGE-M3优化的分片策略，检索效果提升30%+')
    
    doc.add_heading('最佳实践: 基于BGE-M3模型特性的推荐配置', level=4)
    
    # 表格5：分片配置
    table5 = doc.add_table(rows=5, cols=5)
    table5.style = 'Table Grid'
    for i, header in enumerate(['文档类型', 'chunk_size', 'chunk_overlap', '策略', '原因']):
        cell = table5.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E6F3FF')
    
    data5 = [
        ['法律法规', '512字符', '128字符 (25%)', '段落感知', '条款完整性重要'],
        ['产品文档', '768字符', '192字符 (25%)', '固定长度', '内容连续性强'],
        ['案例研究', '1024字符', '256字符 (25%)', '段落感知', '叙事完整性'],
        ['FAQ/问答', '384字符', '96字符 (25%)', '问答对', '一问一答独立']
    ]
    for i, row_data in enumerate(data5):
        for j, cell_text in enumerate(row_data):
            table5.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph('✅ 默认配置: chunk_size=512, overlap=128 - 适用于大多数合规文档')
    
    # ========== 第三部分：实战篇 ==========
    doc.add_heading('第三部分：实战篇', level=1)
    
    doc.add_heading('2.3 综合应用', level=2)
    doc.add_paragraph('进入界面后可在引导下咨询关心的问题。')
    
    # 插入第八张图片
    if image_paths[7]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[7], width=Inches(6))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[7])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('2.4 运营监控体系与持续优化', level=2)
    doc.add_paragraph('接入langfuse，对系统运行情况监控，针对性进行成本、性能分析和优化。')
    doc.add_paragraph('系统具备点赞、点踩功能，针对点踩的badcase可基于后台数据分析改进。')
    
    # 插入第九张图片
    if image_paths[8]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_paths[8], width=Inches(5))
        caption = doc.add_paragraph(IMAGE_DESCRIPTIONS[8])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # ========== 总结 ==========
    doc.add_heading('3.1 总结与后续规划', level=1)
    
    doc.add_heading('核心要点回顾', level=2)
    doc.add_paragraph('• 问题本质：出海合规评审最大挑战是信息缺失、合规评审专业性强和流程环节长')
    doc.add_paragraph('• 解决方案：AI+知识库将历史经验沉淀为智能评估系统')
    doc.add_paragraph('• 核心成果：评估周期↓ X%、风险提前X个月应对、年度节省XXX')
    doc.add_paragraph('• 技术栈：AI Agent+LLM+工作流+30大类知识库+6大核心功能')
    
    doc.add_heading('后续规划', level=2)
    doc.add_paragraph('• 更多知识库：接入更多产品，丰富知识库维度')
    doc.add_paragraph('• 更多客户场景：工作流打磨和接入更多客户场景，方案持续打磨')
    doc.add_paragraph('• 知识库分库：设置知识优先级，工作流根据查询类型路由到不同知识库')
    doc.add_paragraph('• 实时更新：知识库实时自动更新')
    doc.add_paragraph('• 评测与改进：知识库评测')
    
    # 结语
    doc.add_paragraph()
    ending = doc.add_paragraph()
    ending_run = ending.add_run('🌏 让每一个客户和产品都能成功扬帆出海')
    ending_run.font.bold = True
    ending_run.font.size = Pt(14)
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    final_para = doc.add_paragraph()
    final_para.add_run('"智能新航海，合规不触礁"。面临出海和AI的大航海时代，采用AI技术，助力云业务出海合规评估、解决方案，打造出海合规咨询服务，护航云业务和中国企业出海。')
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def main():
    print("=" * 50)
    print("KM文章转Word文档工具")
    print("=" * 50)
    
    # 创建session用于下载
    session = requests.Session()
    session.headers.update({
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Referer': 'https://km.woa.com/'
    })
    
    # 下载图片
    print("\n[1/2] 下载图片...")
    image_paths = download_images(session)
    
    # 创建文档
    print("\n[2/2] 创建Word文档...")
    doc = create_document(image_paths)
    
    # 保存文档
    output_path = os.path.join(WORK_DIR, '扬帆出海_AI赋能云业务全球化合规实战.docx')
    doc.save(output_path)
    print(f"\n✅ 文档已保存: {output_path}")
    print("=" * 50)

if __name__ == '__main__':
    main()
